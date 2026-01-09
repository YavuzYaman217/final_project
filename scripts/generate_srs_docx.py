#!/usr/bin/env python3
"""
Generate SRS Document with Embedded Diagrams as DOCX
This script creates a proper Word document with all diagrams embedded as images.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

# Paths
BASE_DIR = "/home/ubuntu/final_project"
DIAGRAMS_DIR = os.path.join(BASE_DIR, "diagrams")
OUTPUT_FILE = os.path.join(BASE_DIR, "analysis", "SRS_Document_With_Diagrams.docx")

def add_heading(doc, text, level=1):
    """Add a heading to the document."""
    doc.add_heading(text, level=level)

def add_paragraph(doc, text, bold=False):
    """Add a paragraph to the document."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    return p

def add_bullet_list(doc, items):
    """Add a bullet list to the document."""
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

def add_numbered_list(doc, items):
    """Add a numbered list to the document."""
    for item in items:
        doc.add_paragraph(item, style='List Number')

def add_image(doc, image_path, caption, width_inches=6.0):
    """Add an image with caption to the document."""
    if os.path.exists(image_path):
        doc.add_picture(image_path, width=Inches(width_inches))
        # Center the image
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add caption
        caption_para = doc.add_paragraph()
        caption_run = caption_para.add_run(caption)
        caption_run.italic = True
        caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()  # Add spacing
    else:
        doc.add_paragraph(f"[Image not found: {image_path}]")

def create_srs_document():
    """Create the complete SRS document with embedded diagrams."""
    doc = Document()
    
    # ========== TITLE PAGE ==========
    title = doc.add_heading('Software Requirements Specification (SRS)', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run('Exam Security System')
    subtitle_run.bold = True
    subtitle_run.font.size = Pt(18)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    subtitle2 = doc.add_paragraph()
    subtitle2_run = subtitle2.add_run('(Identity Verification + Seating Plan + Violation Logging)')
    subtitle2_run.font.size = Pt(14)
    subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Document info
    info_para = doc.add_paragraph()
    info_para.add_run('Document Version: ').bold = True
    info_para.add_run('2.0 (With Embedded Diagrams)\n')
    info_para.add_run('Date: ').bold = True
    info_para.add_run('January 2026\n')
    info_para.add_run('Project Name: ').bold = True
    info_para.add_run('Exam Security System\n')
    info_para.add_run('Course: ').bold = True
    info_para.add_run('Software Testing & Validation')
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # ========== TABLE OF CONTENTS ==========
    add_heading(doc, 'Table of Contents', 1)
    toc_items = [
        '1. Introduction',
        '2. Overall Description',
        '3. System Architecture & Diagrams',
        '   3.1 Use Case Diagram',
        '   3.2 Entity Relationship Diagram (ERD)',
        '   3.3 Activity Diagram',
        '   3.4 Sequence Diagrams',
        '4. Functional Requirements',
        '5. Non-Functional Requirements',
        '6. Database Requirements',
        '7. Business Rules',
        '8. User Interface Requirements',
        '9. Testing & Validation Requirements',
        '10. Glossary'
    ]
    for item in toc_items:
        doc.add_paragraph(item)
    
    doc.add_page_break()
    
    # ========== 1. INTRODUCTION ==========
    add_heading(doc, '1. Introduction', 1)
    
    add_heading(doc, '1.1 Purpose', 2)
    doc.add_paragraph(
        'This document specifies the functional and non-functional requirements for the Exam Security System, '
        'a web-based application designed to manage exam-day security operations. The system ensures that only '
        'registered students enter the exam room, that they sit in assigned seats according to the seating plan, '
        'and that all violations are properly recorded and reported.'
    )
    
    add_heading(doc, '1.2 Scope', 2)
    doc.add_paragraph('The Exam Security System is a web-based application that supports three core actors:')
    add_bullet_list(doc, [
        'Students: Individuals taking the exam',
        'Proctors (Invigilators): Personnel monitoring exam compliance',
        'Exam Coordinators (Administrators): Personnel managing exams, seating plans, and reports'
    ])
    doc.add_paragraph(
        'The system integrates a simple machine learning/computer vision component for identity verification '
        'and provides comprehensive violation logging and reporting capabilities.'
    )
    
    add_heading(doc, '1.3 Document Conventions', 2)
    add_bullet_list(doc, [
        'Shall/Must: Indicates a mandatory requirement',
        'Should: Indicates a recommended requirement',
        'May: Indicates an optional requirement',
        'FR-X: Functional Requirement identifier',
        'NFR-X: Non-Functional Requirement identifier'
    ])
    
    add_heading(doc, '1.4 Intended Audience', 2)
    add_bullet_list(doc, [
        'Development Team',
        'Quality Assurance Team',
        'Project Stakeholders',
        'Instructors and Evaluators'
    ])
    
    doc.add_page_break()
    
    # ========== 2. OVERALL DESCRIPTION ==========
    add_heading(doc, '2. Overall Description', 1)
    
    add_heading(doc, '2.1 Product Perspective', 2)
    doc.add_paragraph(
        'The Exam Security System is a standalone web application that operates independently but may integrate '
        'with existing student information systems for roster import. It is designed to be deployed in a controlled '
        'exam environment with internet connectivity.'
    )
    
    add_heading(doc, '2.2 Product Functions', 2)
    doc.add_paragraph('The system provides the following major functions:')
    add_numbered_list(doc, [
        'Authentication & Authorization: Role-based access control for Proctors and Administrators',
        'Exam Management: Create and configure exams with date, time, and room information',
        'Seating Plan Management: Define and manage student seating assignments',
        'Student Roster Management: Import or manually enter student information',
        'Identity Verification: Capture and verify student identity using photo comparison',
        'Check-In Workflow: Process student check-in with photo capture and verification',
        'Seat Compliance Verification: Validate that students sit in assigned seats',
        'Violation Recording: Log and document any exam violations',
        'Reporting: Generate reports on check-ins, mismatches, and violations'
    ])
    
    add_heading(doc, '2.3 User Classes and Characteristics', 2)
    
    add_heading(doc, '2.3.1 Students', 3)
    add_bullet_list(doc, [
        'Characteristics: Exam participants, may have limited technical experience',
        'Responsibilities: Provide identity verification, sit in assigned seat',
        'Frequency of Use: One-time per exam session'
    ])
    
    add_heading(doc, '2.3.2 Proctors (Invigilators)', 3)
    add_bullet_list(doc, [
        'Characteristics: Trained exam monitors, moderate technical experience',
        'Responsibilities: Verify student identity, check seating compliance, record violations',
        'Frequency of Use: Throughout exam duration'
    ])
    
    add_heading(doc, '2.3.3 Exam Coordinators (Administrators)', 3)
    add_bullet_list(doc, [
        'Characteristics: Exam management personnel, good technical experience',
        'Responsibilities: Create exams, manage seating plans, import rosters, generate reports',
        'Frequency of Use: Before and after exam sessions'
    ])
    
    add_heading(doc, '2.4 Operating Environment', 2)
    add_bullet_list(doc, [
        'Platform: Web-based application (browser-based)',
        'Browsers: Chrome, Firefox, Safari, Edge (latest versions)',
        'Server: Node.js/Express or Python/Flask backend',
        'Database: MySQL, PostgreSQL, or similar relational database',
        'Hardware: Standard desktop/laptop with camera for photo capture'
    ])
    
    add_heading(doc, '2.5 Design and Implementation Constraints', 2)
    add_bullet_list(doc, [
        'Simple ML/Computer Vision component (library-based, not custom-trained models)',
        'No deep learning model training required',
        'Grading focuses on integration, validation, and workflow correctness',
        'Role-based access control must be enforced',
        'System must handle concurrent user sessions'
    ])
    
    add_heading(doc, '2.6 Assumptions and Dependencies', 2)
    add_bullet_list(doc, [
        'Students have valid registered accounts with the system',
        'Photo capture devices (cameras) are available at check-in stations',
        'Network connectivity is stable during exam sessions',
        'Database is properly backed up and maintained',
        'ML/CV library (e.g., face_recognition, OpenCV) is available'
    ])
    
    doc.add_page_break()
    
    # ========== 3. SYSTEM ARCHITECTURE & DIAGRAMS ==========
    add_heading(doc, '3. System Architecture & Diagrams', 1)
    doc.add_paragraph(
        'This section provides a visual overview of the system architecture through a series of UML diagrams. '
        'All diagrams follow UML 2.5 standards and are rendered at high resolution (300 DPI) for clarity.'
    )
    
    # 3.1 Use Case Diagram
    add_heading(doc, '3.1 Use Case Diagram', 2)
    doc.add_paragraph(
        'The Use Case Diagram illustrates the interactions between system actors (Administrator, Proctor, Student, '
        'and ML Service) and the main use cases of the Exam Security System. It shows the functional scope of the '
        'system from the users\' perspective.'
    )
    add_image(doc, os.path.join(DIAGRAMS_DIR, "usecase.png"), "Figure 1: Use Case Diagram", 6.5)
    
    # 3.2 ERD
    add_heading(doc, '3.2 Entity Relationship Diagram (ERD)', 2)
    doc.add_paragraph(
        'The Entity Relationship Diagram shows the complete database schema, including all entities (tables), '
        'their attributes, and the relationships between them. It uses Crow\'s Foot notation to represent '
        'cardinality. The diagram includes 9 main tables: users, rooms, exams, students, exam_rosters, '
        'seating_plans, seat_assignments, check_ins, violations, and audit_logs.'
    )
    add_image(doc, os.path.join(DIAGRAMS_DIR, "erd_updated.png"), "Figure 2: Entity Relationship Diagram (ERD)", 6.5)
    
    # 3.3 Activity Diagram
    add_heading(doc, '3.3 Activity Diagram', 2)
    doc.add_paragraph(
        'The Activity Diagram provides a high-level overview of the exam day workflow, showing the flow of '
        'activities between the Administrator, Proctor, and the System. It illustrates the complete process '
        'from exam setup to report generation.'
    )
    add_image(doc, os.path.join(DIAGRAMS_DIR, "activity.png"), "Figure 3: Activity Diagram (Exam Day Workflow)", 6.5)
    
    doc.add_page_break()
    
    # 3.4 Sequence Diagrams
    add_heading(doc, '3.4 Sequence Diagrams', 2)
    
    add_heading(doc, '3.4.1 Student Check-In Workflow', 3)
    doc.add_paragraph(
        'This sequence diagram details the step-by-step interactions for the student check-in process, '
        'including identity verification using ML/CV, seat compliance checking, and automatic violation creation. '
        'It shows the communication between the Proctor, Web Interface, Backend Server, ML Service, Database, '
        'and File Storage components.'
    )
    add_image(doc, os.path.join(DIAGRAMS_DIR, "sequence_checkin_detailed.png"), 
              "Figure 4: Sequence Diagram - Student Check-In Workflow", 6.5)
    
    add_heading(doc, '3.4.2 Violation Recording Workflow', 3)
    doc.add_paragraph(
        'This sequence diagram shows the workflow for a proctor recording an exam violation with optional '
        'evidence attachment. It includes violation type selection, severity assignment, and audit logging.'
    )
    add_image(doc, os.path.join(DIAGRAMS_DIR, "sequence_violation.png"), 
              "Figure 5: Sequence Diagram - Violation Recording Workflow", 6.5)
    
    add_heading(doc, '3.4.3 Report Generation Workflow', 3)
    doc.add_paragraph(
        'This sequence diagram illustrates the process of an administrator generating and exporting exam reports. '
        'It shows the interaction with the Report Generator Service and the various export options (CSV, PDF).'
    )
    add_image(doc, os.path.join(DIAGRAMS_DIR, "sequence_reporting.png"), 
              "Figure 6: Sequence Diagram - Report Generation Workflow", 6.5)
    
    doc.add_page_break()
    
    # ========== 4. FUNCTIONAL REQUIREMENTS ==========
    add_heading(doc, '4. Functional Requirements', 1)
    
    # 4.1 Authentication
    add_heading(doc, '4.1 Authentication & Authorization (FR-1 to FR-3)', 2)
    
    add_heading(doc, 'FR-1: User Login', 3)
    add_bullet_list(doc, [
        'Description: Users shall authenticate using username and password',
        'Actors: Proctor, Administrator',
        'Preconditions: User account exists in the system',
        'Steps: 1) User navigates to login page, 2) User enters username and password, 3) System validates credentials against database, 4) System creates session and redirects to dashboard',
        'Postconditions: User is authenticated and session is active',
        'Alternative Flows: Invalid credentials trigger error message; account lockout after 5 failed attempts'
    ])
    
    add_heading(doc, 'FR-2: Role-Based Access Control', 3)
    add_bullet_list(doc, [
        'Description: System shall enforce role-based access control for Proctor and Administrator roles',
        'Actors: System',
        'Rules: Proctors can only access check-in, violation recording, and basic reporting features; Administrators can access all features including exam creation, roster management, and advanced reporting; Students do not require login',
        'Validation: Unauthorized access attempts shall be logged and rejected'
    ])
    
    add_heading(doc, 'FR-3: Session Management', 3)
    add_bullet_list(doc, [
        'Description: System shall manage user sessions with automatic timeout',
        'Timeout Duration: 30 minutes of inactivity',
        'Actions: Expired sessions redirect users to login page',
        'Validation: Session tokens are validated on each request'
    ])
    
    # 4.2 Exam Management
    add_heading(doc, '4.2 Exam Management (FR-4 to FR-6)', 2)
    
    add_heading(doc, 'FR-4: Exam Creation', 3)
    add_bullet_list(doc, [
        'Description: Administrators shall create exams with exam name/code, date and time, room/location, duration, and maximum capacity',
        'Validation Rules: Exam code must be unique, date/time must be in the future, capacity must be greater than 0',
        'Postconditions: Exam is created and available for roster and seating plan assignment'
    ])
    
    add_heading(doc, 'FR-5: Exam Configuration', 3)
    add_bullet_list(doc, [
        'Description: Administrators shall configure exam parameters: enable/disable identity verification requirement, set seating plan requirement (mandatory/optional), configure violation categories and severity levels',
        'Validation: Configuration changes are logged with timestamp and user ID'
    ])
    
    add_heading(doc, 'FR-6: Exam Status Management', 3)
    add_bullet_list(doc, [
        'Description: System shall track exam status: Draft, Active, Completed, Archived',
        'Transitions: Only authorized transitions are allowed (Draft → Active → Completed → Archived)'
    ])
    
    # 4.3 Student Roster Management
    add_heading(doc, '4.3 Student Roster Management (FR-7 to FR-9)', 2)
    
    add_heading(doc, 'FR-7: Student Roster Import', 3)
    add_bullet_list(doc, [
        'Description: Administrators shall import student roster via CSV file upload',
        'File Format: CSV with columns: StudentID, FirstName, LastName, Email, RegistrationNumber',
        'Validation Rules: StudentID must be unique, required fields must not be empty, duplicate entries are flagged for review',
        'Postconditions: Students are added to the exam roster'
    ])
    
    add_heading(doc, 'FR-8: Manual Student Entry', 3)
    add_bullet_list(doc, [
        'Description: Administrators shall manually add individual students to the roster',
        'Fields Required: StudentID, FirstName, LastName, Email, RegistrationNumber',
        'Validation: Same rules as FR-7',
        'Postconditions: Student is added to roster'
    ])
    
    add_heading(doc, 'FR-9: Roster Management', 3)
    add_bullet_list(doc, [
        'Description: Administrators shall manage the student roster: view all students, edit student information, remove students, export roster to CSV',
        'Validation: Changes are logged with timestamp and user ID'
    ])
    
    # 4.4 Seating Plan Management
    add_heading(doc, '4.4 Seating Plan Management (FR-10 to FR-12)', 2)
    
    add_heading(doc, 'FR-10: Seating Plan Creation', 3)
    add_bullet_list(doc, [
        'Description: Administrators shall create seating plans with Grid-Based (rows and columns) or Seat Code-Based (individual seat codes) options',
        'Validation Rules: Total seats must be ≥ number of students in roster, seat identifiers must be unique, plan must be associated with an exam',
        'Postconditions: Seating plan is created and ready for student assignment'
    ])
    
    add_heading(doc, 'FR-11: Student Seat Assignment', 3)
    add_bullet_list(doc, [
        'Description: Administrators shall assign students to seats via manual assignment (drag-and-drop or form-based) or automatic assignment (random or sequential)',
        'Validation Rules: Each student assigned to exactly one seat, each seat assigned to at most one student, cannot assign students not in roster',
        'Postconditions: All students have assigned seats'
    ])
    
    add_heading(doc, 'FR-12: Seating Plan Visualization', 3)
    add_bullet_list(doc, [
        'Description: System shall display seating plan with visual grid representation, student names/IDs in assigned seats, color coding for assigned/unassigned/occupied seats, real-time updates during check-in',
        'Actors: Proctor, Administrator',
        'Validation: Display updates within 5 seconds of check-in'
    ])
    
    # 4.5 Identity Verification
    add_heading(doc, '4.5 Identity Verification (FR-13 to FR-15)', 2)
    
    add_heading(doc, 'FR-13: Photo Capture', 3)
    add_bullet_list(doc, [
        'Description: System shall capture student photo during check-in',
        'Process: 1) Student positions face in front of camera, 2) System captures image automatically or on manual trigger, 3) Image is stored with timestamp and student ID',
        'Technical Requirements: Support multiple image formats (JPEG, PNG), image resolution ≥ 640×480 pixels, automatic face detection to guide student positioning',
        'Validation: Image must contain a detectable face'
    ])
    
    add_heading(doc, 'FR-14: ML/Computer Vision Verification', 3)
    add_bullet_list(doc, [
        'Description: System shall verify captured photo against registered student photo using ML/CV',
        'Implementation Options: Face verification using face_recognition library, ID photo similarity comparison using OpenCV, basic template matching with embeddings',
        'Process: 1) Extract face embeddings from captured photo, 2) Compare with registered student photo embeddings, 3) Generate match confidence score (0-100%), 4) Determine pass/fail based on threshold (e.g., 75%)',
        'Output: Match result (Match/No Match) with confidence score',
        'Validation: Threshold must be configurable, results must be logged with timestamp, manual override available for Proctors'
    ])
    
    add_heading(doc, 'FR-15: Verification Decision', 3)
    add_bullet_list(doc, [
        'Description: System shall present verification decision to Proctor',
        'Display Information: Captured photo, registered student photo, match confidence score, recommendation (Match/No Match)',
        'Proctor Actions: Accept verification, reject verification, override decision (manual approval despite mismatch)',
        'Postconditions: Verification result is recorded with timestamp and Proctor ID'
    ])
    
    # 4.6 Check-In Workflow
    add_heading(doc, '4.6 Check-In Workflow (FR-16 to FR-20)', 2)
    
    add_heading(doc, 'FR-16: Check-In Initiation', 3)
    add_bullet_list(doc, [
        'Description: Proctor initiates check-in process for a student',
        'Input: Student ID or name search',
        'Process: 1) System retrieves student information from roster, 2) System displays student details and assigned seat, 3) System prompts for photo capture',
        'Validation: Student must be in roster and not already checked in'
    ])
    
    add_heading(doc, 'FR-17: Photo Capture & Upload', 3)
    add_bullet_list(doc, [
        'Description: System captures and uploads student photo during check-in',
        'Process: 1) Camera interface is displayed, 2) Student positions face in frame, 3) Photo is captured, 4) Photo is uploaded to server, 5) System confirms successful upload',
        'Validation: Photo must be valid and contain detectable face'
    ])
    
    add_heading(doc, 'FR-18: ML Verification Decision', 3)
    add_bullet_list(doc, [
        'Description: System processes photo through ML/CV component',
        'Process: 1) System extracts face embeddings, 2) System compares with registered student photo, 3) System generates confidence score, 4) System presents result to Proctor',
        'Output: Match/No Match with confidence score',
        'Validation: Result is logged with timestamp'
    ])
    
    add_heading(doc, 'FR-19: Seat Compliance Check', 3)
    add_bullet_list(doc, [
        'Description: System verifies that student is sitting in assigned seat',
        'Process: 1) Proctor confirms student identity, 2) Proctor verifies student is in correct seat, 3) System records seat assignment and check-in time',
        'Validation: Seat must match student\'s assigned seat',
        'Alternative: If student is in wrong seat, violation is recorded'
    ])
    
    add_heading(doc, 'FR-20: Check-In Completion', 3)
    add_bullet_list(doc, [
        'Description: System completes check-in process and records result',
        'Recorded Information: Student ID, check-in timestamp, verification result, assigned seat, actual seat, proctor ID, any violations or notes',
        'Postconditions: Student is marked as checked-in; seating plan is updated'
    ])
    
    # 4.7 Violation Recording
    add_heading(doc, '4.7 Violation Recording (FR-21 to FR-24)', 2)
    
    add_heading(doc, 'FR-21: Violation Categories', 3)
    add_bullet_list(doc, [
        'Description: System shall support violation categories: Identity Mismatch, Seat Mismatch, Unauthorized Materials, Disruptive Behavior, Late Arrival, Other',
        'Validation: Each violation must have a category'
    ])
    
    add_heading(doc, 'FR-22: Violation Recording', 3)
    add_bullet_list(doc, [
        'Description: Proctor shall record violations with violation category, student ID, timestamp, reason/notes, evidence image (optional), severity level (Low/Medium/High)',
        'Validation Rules: All required fields must be completed, timestamp must be during exam session, evidence image must be valid image file',
        'Postconditions: Violation is recorded and associated with student'
    ])
    
    add_heading(doc, 'FR-23: Violation Evidence', 3)
    add_bullet_list(doc, [
        'Description: Proctor may attach evidence image to violation record',
        'Process: 1) Proctor captures or uploads image as evidence, 2) System stores image with violation record, 3) Image is linked to violation ID',
        'Validation: Image must be valid format (JPEG, PNG)',
        'Optional: Evidence images are optional but recommended'
    ])
    
    add_heading(doc, 'FR-24: Violation Status Tracking', 3)
    add_bullet_list(doc, [
        'Description: System shall track violation status: Recorded, Reviewed, Resolved, Dismissed',
        'Transitions: Only authorized transitions allowed',
        'Validation: Status changes are logged with timestamp and user ID'
    ])
    
    # 4.8 Reporting
    add_heading(doc, '4.8 Reporting (FR-25 to FR-28)', 2)
    
    add_heading(doc, 'FR-25: Check-In Report', 3)
    add_bullet_list(doc, [
        'Description: System shall generate check-in report with list of all students checked in, check-in timestamp, verification result, assigned vs. actual seat, proctor who performed check-in',
        'Format: Exportable to CSV, PDF, or display in web interface',
        'Filters: By exam, by date range, by proctor'
    ])
    
    add_heading(doc, 'FR-26: Mismatch Report', 3)
    add_bullet_list(doc, [
        'Description: System shall generate mismatch report with students with identity mismatches, students in wrong seats, timestamp of mismatch detection, proctor who recorded mismatch',
        'Format: Exportable to CSV, PDF',
        'Filters: By exam, by confidence score range'
    ])
    
    add_heading(doc, 'FR-27: Violation Report', 3)
    add_bullet_list(doc, [
        'Description: System shall generate violation report with all violations recorded, filterable by category, severity, status, student, proctor, includes violation details and evidence links',
        'Format: Exportable to CSV, PDF',
        'Filters: By exam, by date, by violation type'
    ])
    
    add_heading(doc, 'FR-28: Summary Report', 3)
    add_bullet_list(doc, [
        'Description: System shall generate summary report with total students checked in, total identity mismatches, total seat mismatches, total violations by category, check-in compliance percentage',
        'Format: Dashboard view with charts and graphs',
        'Filters: By exam'
    ])
    
    doc.add_page_break()
    
    # ========== 5. NON-FUNCTIONAL REQUIREMENTS ==========
    add_heading(doc, '5. Non-Functional Requirements', 1)
    
    add_heading(doc, '5.1 Performance (NFR-1 to NFR-3)', 2)
    add_bullet_list(doc, [
        'NFR-1: Response Time - All API responses shall be < 1 second under normal load',
        'NFR-2: Verification Speed - ML/CV verification shall complete in < 3 seconds',
        'NFR-3: Concurrent Users - System shall support at least 10 concurrent proctors'
    ])
    
    add_heading(doc, '5.2 Security (NFR-4 to NFR-6)', 2)
    add_bullet_list(doc, [
        'NFR-4: Password Hashing - All user passwords shall be hashed using bcrypt',
        'NFR-5: Data Encryption - All data in transit shall be encrypted using TLS/SSL',
        'NFR-6: Audit Trail - All critical actions shall be logged in an immutable audit trail'
    ])
    
    add_heading(doc, '5.3 Usability (NFR-7 to NFR-8)', 2)
    add_bullet_list(doc, [
        'NFR-7: User Interface - UI shall be intuitive and require minimal training',
        'NFR-8: Accessibility - System shall comply with WCAG 2.1 AA standards'
    ])
    
    add_heading(doc, '5.4 Reliability (NFR-9 to NFR-10)', 2)
    add_bullet_list(doc, [
        'NFR-9: Uptime - System shall have 99.9% uptime during exam periods',
        'NFR-10: Data Integrity - Database shall enforce referential integrity through foreign keys'
    ])
    
    doc.add_page_break()
    
    # ========== 6. DATABASE REQUIREMENTS ==========
    add_heading(doc, '6. Database Requirements', 1)
    
    add_heading(doc, '6.1 Database Schema', 2)
    doc.add_paragraph(
        'The database schema is detailed in the Entity Relationship Diagram (ERD) in Section 3.2. '
        'It consists of 10 main tables:'
    )
    add_numbered_list(doc, [
        'users: System users (Admins and Proctors)',
        'rooms: Exam room information',
        'exams: Exam configurations',
        'students: Student information and registered photos',
        'exam_rosters: Student-to-exam assignments',
        'seating_plans: Seating plan configurations',
        'seat_assignments: Individual seat records',
        'check_ins: Check-in records with verification results',
        'violations: Violation records with evidence',
        'audit_logs: System activity audit trail'
    ])
    
    doc.add_page_break()
    
    # ========== 7. BUSINESS RULES ==========
    add_heading(doc, '7. Business Rules', 1)
    doc.add_paragraph(
        'This section outlines the key business rules that govern the behavior of the Exam Security System. '
        'These rules ensure data integrity, security, and consistent operational workflows.'
    )
    
    add_heading(doc, 'BR-1: Exam Status Progression', 2)
    add_bullet_list(doc, [
        'Description: An exam must follow a specific lifecycle: DRAFT → ACTIVE → COMPLETED → ARCHIVED',
        'Rationale: Ensures that exams are properly configured before becoming active and are properly closed after completion',
        'Validation: The system will only allow status transitions in the specified order',
        'Error Handling: An error message will be displayed if an invalid status transition is attempted',
        'Related Requirements: FR-4, FR-5'
    ])
    
    add_heading(doc, 'BR-2: Proctor-Exam Assignment', 2)
    add_bullet_list(doc, [
        'Description: A proctor must be assigned to an exam to perform check-ins and record violations for that exam',
        'Rationale: Enforces accountability and ensures that only authorized proctors can manage a specific exam session',
        'Validation: The system will check for a valid proctor-exam assignment before allowing any operational actions',
        'Error Handling: Access will be denied with a "Not authorized for this exam" message',
        'Related Requirements: FR-2, FR-15'
    ])
    
    add_heading(doc, 'BR-3: Identity Verification Threshold', 2)
    add_bullet_list(doc, [
        'Description: A confidence score of 75% or higher from the ML/CV service is required for an automatic identity match',
        'Rationale: Balances security with usability by setting a reasonable threshold for automated verification',
        'Validation: The system will check if confidence_score >= 0.75 to determine the verification result',
        'Error Handling: Scores below 75% will be flagged as a "NO_MATCH", requiring manual review',
        'Related Requirements: FR-14, FR-19'
    ])
    
    add_heading(doc, 'BR-4: Account Lockout Policy', 2)
    add_bullet_list(doc, [
        'Description: A user account will be temporarily locked after 5 consecutive failed login attempts',
        'Rationale: Prevents brute-force attacks on user accounts',
        'Validation: The system will track the number of failed login attempts for each user',
        'Error Handling: After 5 failed attempts, the user account will be deactivated',
        'Related Requirements: FR-1'
    ])
    
    add_heading(doc, 'BR-5: Violation-Check-In Association', 2)
    add_bullet_list(doc, [
        'Description: All recorded violations must be linked to a valid check-in record',
        'Rationale: Ensures traceability and context for every violation',
        'Validation: A foreign key constraint (check_in_id) will enforce this relationship',
        'Error Handling: The system will prevent the creation of a violation without a valid check-in',
        'Related Requirements: FR-21'
    ])
    
    add_heading(doc, 'BR-6: Seat Assignment Timing', 2)
    add_bullet_list(doc, [
        'Description: Seat assignments for an exam can only be created or modified before the exam start time',
        'Rationale: Prevents changes to the seating plan while an exam is in progress',
        'Validation: The system will check if current_time < exam.start_time before allowing modifications',
        'Error Handling: An error message will be displayed if modification is attempted during active exam',
        'Related Requirements: FR-11'
    ])
    
    add_heading(doc, 'BR-7: Photo Requirement for Verification', 2)
    add_bullet_list(doc, [
        'Description: A registered student photo is mandatory for the ML/CV identity verification process',
        'Rationale: The system cannot perform a comparison without a reference photo',
        'Validation: The system will check if the registered_photo_path is not NULL or empty',
        'Error Handling: If no registered photo exists, manual proctor approval is required',
        'Related Requirements: FR-13, FR-16'
    ])
    
    add_heading(doc, 'BR-8: Role-Based Report Access', 2)
    add_bullet_list(doc, [
        'Description: Proctors can only view reports for exams they are assigned to, while Administrators can view all reports',
        'Rationale: Enforces data privacy and the principle of least privilege',
        'Validation: The system will filter report data based on the user\'s role and exam assignments',
        'Error Handling: An HTTP 403 Forbidden error will be returned for unauthorized access attempts',
        'Related Requirements: FR-2, FR-25'
    ])
    
    doc.add_page_break()
    
    # ========== 8. USER INTERFACE REQUIREMENTS ==========
    add_heading(doc, '8. User Interface Requirements', 1)
    doc.add_paragraph(
        'The user interface shall be designed with the following principles in mind:'
    )
    add_bullet_list(doc, [
        'Intuitive navigation with clear menu structure',
        'Responsive design for desktop and tablet devices',
        'Consistent color coding for status indicators',
        'Real-time feedback for user actions',
        'Accessible design following WCAG 2.1 AA guidelines'
    ])
    
    doc.add_page_break()
    
    # ========== 9. TESTING & VALIDATION REQUIREMENTS ==========
    add_heading(doc, '9. Testing & Validation Requirements', 1)
    doc.add_paragraph(
        'The system shall undergo comprehensive testing including:'
    )
    add_bullet_list(doc, [
        'Unit Testing: All individual functions and methods',
        'Integration Testing: API endpoints and database operations',
        'System Testing: End-to-end workflows',
        'Performance Testing: Load and stress testing',
        'Security Testing: Authentication and authorization',
        'User Acceptance Testing: Real-world scenario validation'
    ])
    doc.add_paragraph(
        'Refer to the test-docs/test_cases.md document for detailed test cases.'
    )
    
    doc.add_page_break()
    
    # ========== 10. GLOSSARY ==========
    add_heading(doc, '10. Glossary', 1)
    
    # Create a table for glossary
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    
    # Header row
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Term'
    header_cells[1].text = 'Definition'
    
    # Make header bold
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    # Glossary terms
    glossary = [
        ('Proctor', 'An individual who monitors an exam (also known as an Invigilator)'),
        ('Roster', 'A list of students registered for an exam'),
        ('Seating Plan', 'A diagram showing where each student is assigned to sit'),
        ('SSIM', 'Structural Similarity Index, an algorithm for measuring image similarity'),
        ('Violation', 'Any breach of exam rules'),
        ('Check-In', 'The process of verifying a student\'s identity and recording their presence'),
        ('ML/CV', 'Machine Learning / Computer Vision - technologies used for automated identity verification'),
        ('ERD', 'Entity Relationship Diagram - a visual representation of database structure'),
        ('JWT', 'JSON Web Token - a standard for secure authentication tokens'),
        ('RBAC', 'Role-Based Access Control - a method of restricting system access based on user roles')
    ]
    
    for term, definition in glossary:
        row_cells = table.add_row().cells
        row_cells[0].text = term
        row_cells[1].text = definition
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Footer
    footer_para = doc.add_paragraph()
    footer_para.add_run('--- End of Document ---').italic = True
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Save the document
    doc.save(OUTPUT_FILE)
    print(f"Document saved to: {OUTPUT_FILE}")
    
    # Verify images are embedded
    print("\nVerifying embedded images...")
    from zipfile import ZipFile
    with ZipFile(OUTPUT_FILE, 'r') as zip_file:
        media_files = [f for f in zip_file.namelist() if f.startswith('word/media/')]
        print(f"Found {len(media_files)} embedded images:")
        for f in media_files:
            print(f"  - {f}")

if __name__ == "__main__":
    create_srs_document()
